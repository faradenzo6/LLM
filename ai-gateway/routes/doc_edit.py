from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import StreamingResponse, JSONResponse
from pathlib import Path
from hashlib import sha256
from typing import List, Dict, Any, Tuple
import base64, io, json, re, difflib, os, requests

# PDF
from weasyprint import HTML

router = APIRouter(prefix="/doc", tags=["doc"])

TEMPLATES_DIR = Path("/app/doc-templates")
AUDIT_PATH = "/var/log/ai-gateway/audit.jsonl"

VLLM_URL = os.getenv("VLLM_URL", "http://vllm:8000/v1/chat/completions")
VLLM_API_KEY = os.getenv("VLLM_API_KEY", "")
VLLM_MODEL = os.getenv("VLLM_MODEL", "Qwen/Qwen2-7B-Instruct")


# ------------------------------ utils / audit ------------------------------

def _audit_write(rec: Dict[str, Any]) -> None:
    try:
        os.makedirs(os.path.dirname(AUDIT_PATH), exist_ok=True)
        with open(AUDIT_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    except Exception:
        pass


def _ext(fname: str) -> str:
    return Path(fname).suffix.lower()


def _content_type_by_ext(ext: str) -> str:
    return {
        ".html": "text/html; charset=utf-8",
        ".htm": "text/html; charset=utf-8",
        ".txt": "text/plain; charset=utf-8",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf": "application/pdf",
    }.get(ext, "application/octet-stream")


# ------------------------------ templates API ------------------------------

@router.get("/templates")
async def list_templates():
    items = []
    for f in TEMPLATES_DIR.glob("*.html"):
        text = f.read_text(encoding="utf-8", errors="ignore")
        placeholders = re.findall(r"\{\{(.*?)\}\}", text)
        items.append({
            "id": f.stem,
            "name": f.stem,
            "engine": "simple",
            "placeholders": placeholders
        })
    return items


@router.post("/preview")
async def preview_doc(req: Request):
    data = await req.json()
    tpl_id = data.get("template_id")
    payload = data.get("payload", {})

    tpl_path = TEMPLATES_DIR / f"{tpl_id}.html"
    if not tpl_path.exists():
        raise HTTPException(status_code=404, detail=f"Template {tpl_id} not found")

    text = tpl_path.read_text(encoding="utf-8")
    for k, v in payload.items():
        text = text.replace(f"{{{{{k}}}}}", str(v))

    return {"html": text}


@router.post("/generate")
async def generate_doc(req: Request):
    data = await req.json()
    tpl_id = data.get("template_id")
    payload = data.get("payload", {})
    output = data.get("output", "html")
    filename = data.get("filename", "document")

    tpl_path = TEMPLATES_DIR / f"{tpl_id}.html"
    if not tpl_path.exists():
        raise HTTPException(status_code=404, detail=f"Template {tpl_id} not found")

    text = tpl_path.read_text(encoding="utf-8")
    for k, v in payload.items():
        text = text.replace(f"{{{{{k}}}}}", str(v))

    if output == "html":
        return StreamingResponse(
            io.BytesIO(text.encode("utf-8")),
            media_type="text/html",
            headers={"Content-Disposition": f"attachment; filename={filename}.html"}
        )

    if output == "pdf":
        pdf = HTML(string=text).write_pdf()
        return StreamingResponse(
            io.BytesIO(pdf),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
        )

    raise HTTPException(status_code=400, detail="Invalid output type")


# ------------------------------ edit helpers ------------------------------

def _apply_ops_to_text(text: str, ops: List[Dict[str, Any]]):
    """
    Поддержка:
      - {"action":"replace","pattern":"...","replacement":"...","use_regex":true|false,"count":0|N}
      - {"action":"append_end","text":"..."}   # только для txt/html
      - {"action":"insert_after","pattern":"...","text":"...","use_regex":true|false} # txt/html
    Возвращает: (new_text, stats_list)
    """
    out = text
    stats = []

    for op in ops or []:
        action = op.get("action", "replace")
        use_regex = bool(op.get("use_regex", False))
        count = int(op.get("count", 0) or 0)

        if action == "append_end":
            to_add = op.get("text") or op.get("replacement") or ""
            # гарантируем перевод строки
            if not out.endswith("\n"):
                out += "\n"
            out += to_add
            stats.append({"action": action, "count_done": 1})
            continue

        if action == "insert_after":
            pattern = str(op.get("pattern", ""))
            to_add = op.get("text") or op.get("replacement") or ""
            if use_regex:
                m = re.search(pattern, out, flags=re.DOTALL)
                if m:
                    pos = m.end()
                    out = out[:pos] + ("\n" + to_add) + out[pos:]
                    stats.append({"action": action, "count_done": 1})
                else:
                    stats.append({"action": action, "count_done": 0})
            else:
                idx = out.find(pattern)
                if idx >= 0:
                    pos = idx + len(pattern)
                    out = out[:pos] + ("\n" + to_add) + out[pos:]
                    stats.append({"action": action, "count_done": 1})
                else:
                    stats.append({"action": action, "count_done": 0})
            continue

        # default: replace
        pattern = str(op.get("pattern", ""))
        repl = str(op.get("replacement", ""))
        try:
            if use_regex:
                new_out, n = re.subn(pattern, repl, out, count=0 if count == 0 else count, flags=re.DOTALL)
            else:
                new_out, n = re.subn(re.escape(pattern), repl, out, count=0 if count == 0 else count, flags=re.DOTALL)
        except re.error as e:
            raise HTTPException(status_code=400, detail=f"Bad regex: {e}")
        stats.append({"action": "replace", "pattern": pattern, "count_requested": count, "count_done": n})
        out = new_out

    return out, stats


def _docx_edit(data: bytes, ops: List[Dict[str, Any]]) -> Tuple[bytes, str, str]:
    """
    DOCX:
      - replace: для каждого абзаца;
      - append_paragraph: добавить новый абзац в конец документа;
      - insert_after: вставить абзац после первого подходящего абзаца.
    """
    try:
        from docx import Document
    except Exception:
        raise HTTPException(status_code=500, detail="python-docx not available")

    bio = io.BytesIO(data)
    doc = Document(bio)

    def all_paragraphs(d):
        for p in d.paragraphs:
            yield p
        for t in d.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        yield p

    def collect_text(d) -> str:
        return "\n".join(p.text or "" for p in all_paragraphs(d))

    # 1) replace по параграфам
    replace_ops = [o for o in (ops or []) if o.get("action", "replace") == "replace"]
    for p in all_paragraphs(doc):
        t = p.text
        nt, _ = _apply_ops_to_text(t, replace_ops)
        if nt != t:
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = nt
            else:
                p.add_run(nt)

    # 2) insert_after
    insert_ops = [o for o in (ops or []) if o.get("action") == "insert_after"]
    if insert_ops:
        for op in insert_ops:
            patt = str(op.get("pattern", ""))
            txt = op.get("text") or op.get("replacement") or ""
            use_regex = bool(op.get("use_regex", False))
            done = False
            for p in all_paragraphs(doc):
                hay = p.text or ""
                ok = False
                if use_regex:
                    try:
                        ok = re.search(patt, hay) is not None
                    except re.error:
                        ok = False
                else:
                    ok = patt in hay
                if ok:
                    try:
                        p.insert_paragraph_after(txt)
                    except Exception:
                        # fallback: в конец
                        doc.add_paragraph(txt)
                    done = True
                    break
            if not done:
                # если не нашли якорь — просто в конец
                doc.add_paragraph(txt)

    # 3) append_paragraph
    append_ops = [o for o in (ops or []) if o.get("action") in ("append_paragraph", "append_end")]
    for op in append_ops:
        txt = op.get("text") or op.get("replacement") or ""
        doc.add_paragraph(txt)

    before = ""  # для больших docx diff не обязателен, но соберём после
    after = collect_text(doc)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    diff = ""  # при желании можно формировать diff между before/after
    return out.getvalue(), after, diff


def _peek_text_for_llm(filename: str, data: bytes, limit: int = 0) -> str:
    ext = _ext(filename or "")
    if ext in (".html", ".htm", ".txt", ""):
        s = data.decode("utf-8", errors="replace")
        return s if limit == 0 else s[:limit]
    if ext == ".docx":
        try:
            from docx import Document
            bio = io.BytesIO(data)
            doc = Document(bio)
            chunks = []
            for p in doc.paragraphs:
                if p.text:
                    chunks.append(p.text)
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            if p.text:
                                chunks.append(p.text)
            s = "\n".join(chunks)
            return s if limit == 0 else s[:limit]
        except Exception:
            return ""
    return ""


def _ai_build_ops(instruction: str, text_preview: str) -> List[Dict[str, Any]]:
    if not instruction or len(instruction.strip()) < 3:
        raise HTTPException(status_code=400, detail="Empty instruction")

    sys_prompt = (
        "Ты помощник-редактор документов. Верни ТОЛЬКО JSON (массив операций), без текста вокруг. "
        "Допустимые операции:\n"
        '1) {"action":"replace","pattern":"...","replacement":"...","use_regex":false,"count":0}\n'
        '2) {"action":"append_paragraph","text":"..."}  # для DOCX добавь абзац в конец\n'
        '3) {"action":"insert_after","pattern":"...","text":"...","use_regex":false}\n'
        '4) {"action":"append_end","text":"..."}        # для TXT/HTML допиши в конец\n'
        "count=0 означает заменять везде. Возвращай строго JSON-массив операций."
    )

    user_prompt = (
        f"Инструкция:\n{instruction}\n\n"
        f"Фрагмент документа (возможно обрезан):\n{text_preview}\n\n"
        "Пример простых операций:\n"
        '[{"action":"replace","pattern":"Иванов","replacement":"Петров","use_regex":false,"count":0}]\n'
    )

    payload = {
        "model": VLLM_MODEL,
        "messages": [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.1,
        "max_tokens": 512,
    }
    headers = {"Content-Type": "application/json"}
    if VLLM_API_KEY:
        headers["Authorization"] = f"Bearer {VLLM_API_KEY}"

    r = requests.post(VLLM_URL, headers=headers, json=payload, timeout=60)
    r.raise_for_status()
    content = r.json()["choices"][0]["message"]["content"]

    m = re.search(r"\[.*\]", content, re.S)
    raw = m.group(0) if m else content
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, dict) and "ops" in parsed:
            ops = parsed["ops"]
        elif isinstance(parsed, list):
            ops = parsed
        else:
            raise ValueError("not a list/dict")
    except Exception:
        raise HTTPException(status_code=500, detail="LLM returned invalid JSON for ops")

    if len(ops) > 50:
        ops = ops[:50]
    return ops


# ------------------------------ endpoints: manual edit ------------------------------

@router.post("/edit")
async def doc_edit(request: Request, file: UploadFile = File(...), ops: str = Form("[]")):
    try:
        ops_list = json.loads(ops) if isinstance(ops, str) else ops
    except Exception:
        raise HTTPException(status_code=400, detail="ops must be JSON array")

    data = await file.read()
    ext = _ext(file.filename or "")
    user = request.headers.get("X-User") or "unknown"
    edited = b""
    content_type = "application/octet-stream"
    stats = []

    try:
        if ext in (".html", ".htm", ".txt", ""):
            text = data.decode("utf-8", errors="replace")
            after, stats = _apply_ops_to_text(text, ops_list)
            edited = after.encode("utf-8")
            content_type = _content_type_by_ext(ext or ".txt")
        elif ext == ".docx":
            edited, _, _ = _docx_edit(data, ops_list)
            content_type = _content_type_by_ext(ext)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported extension: {ext}")
    finally:
        pass

    sha = sha256(edited).hexdigest()
    _audit_write({
        "ts": __import__("datetime").datetime.utcnow().isoformat() + "Z",
        "endpoint": "/doc/edit", "ok": True,
        "ip": request.client.host if request.client else None,
        "user": user, "filename": file.filename,
        "edited": (Path(file.filename).stem + ".edited" + ext),
        "ext": ext, "ops_count": len(ops_list),
        "bytes_in": len(data), "bytes_out": len(edited),
        "sha256": sha, "ops_stats": stats,
    })

    resp = {
        "ok": True,
        "filename": (Path(file.filename).stem + ".edited" + ext),
        "content_type": content_type,
        "bytes_in": len(data),
        "bytes_out": len(edited),
        "sha256": sha,
        "diff": "" if ext == ".docx" else "".join(
            difflib.unified_diff(
                data.decode("utf-8", "replace").splitlines(True),
                edited.decode("utf-8", "replace").splitlines(True),
                fromfile=file.filename,
                tofile=(Path(file.filename).stem + ".edited" + ext))
        ),
        "file_base64": base64.b64encode(edited).decode("ascii"),
        "stats": stats,
    }
    return JSONResponse(content=resp, media_type="application/json; charset=utf-8")


# ------------------------------ endpoints: AI edit ------------------------------

@router.post("/ai-ops")
async def doc_ai_ops(request: Request, file: UploadFile = File(...), instruction: str = Form(...)):
    data = await file.read()
    preview = _peek_text_for_llm(file.filename, data)
    ops = _ai_build_ops(instruction, preview)
    _audit_write({
        "ts": __import__("datetime").datetime.utcnow().isoformat() + "Z",
        "endpoint": "/doc/ai-ops",
        "ok": True,
        "filename": file.filename,
        "preview_len": len(preview),
        "ops": ops,
    })
    return JSONResponse(content={"ok": True, "ops": ops}, media_type="application/json; charset=utf-8")


@router.post("/ai-edit")
async def doc_ai_edit(request: Request, file: UploadFile = File(...), instruction: str = Form(...)):
    data = await file.read()
    preview = _peek_text_for_llm(file.filename, data)
    ops = _ai_build_ops(instruction, preview)

    ext = _ext(file.filename or "")
    stats = []
    if ext in (".html", ".htm", ".txt", ""):
        before = data.decode("utf-8", errors="replace")
        after, stats = _apply_ops_to_text(before, ops)
        edited = after.encode("utf-8")
        content_type = _content_type_by_ext(ext or ".txt")
        diff = "".join(difflib.unified_diff(
            before.splitlines(True), after.splitlines(True),
            fromfile=file.filename, tofile=(f"{Path(file.filename).stem}.edited{ext}")
        ))
    elif ext == ".docx":
        edited, after_text, diff = _docx_edit(data, ops)
        content_type = _content_type_by_ext(ext)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported extension: {ext}")

    b64 = base64.b64encode(edited).decode("ascii")
    out_name = f"{Path(file.filename).stem}.edited{ext}"
    result = {
        "ok": True,
        "ops": ops,
        "filename": out_name,
        "content_type": content_type,
        "file_base64": b64,
        "diff": diff,
        "stats": stats,  # для txt/html
    }

    _audit_write({
        "ts": __import__("datetime").datetime.utcnow().isoformat() + "Z",
        "endpoint": "/doc/ai-edit", "ok": True,
        "filename": file.filename,
        "ops": ops, "stats": stats, "diff_len": len(diff),
    })
    return JSONResponse(content=result, media_type="application/json; charset=utf-8")
